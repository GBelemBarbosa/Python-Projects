import docx
import pickle
import re
import os

def extract_from_docx(filename):
    doc = docx.Document(filename)
    # The docx formatting might use bold or something else.
    # Let's extract by looking at bold run or paragraph heading.
    # Another way: just print the first 10 paragraphs so we can see literally what is the text.
    pass

if __name__ == "__main__":
    for i, p in enumerate(docx.Document('Cantrips.docx').paragraphs[:20]):
        print(repr(p.text))
